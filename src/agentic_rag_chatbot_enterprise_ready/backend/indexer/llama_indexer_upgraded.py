"""Production-oriented LlamaIndex document ingestion and retrieval.

This module preserves the original public helpers while migrating the
implementation to current LlamaIndex APIs:

- VectorStoreIndex replaces the legacy GPTVectorStoreIndex alias.
- Global Settings replaces ServiceContext.
- AzureAISearchVectorStore is configured through the current client-based API.
- SHA-256 based local metadata prevents unnecessary re-indexing.
- Deterministic document/chunk IDs make repeated ingestion idempotent.
- PDF resources are explicitly closed.
- Chunking validates its invariants instead of silently accepting bad values.
- Query ``top_k`` is actually applied to retrieval.
"""

from __future__ import annotations

import hashlib
import json
import mimetypes
import os
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

import fitz
import pandas as pd
from azure.core.credentials import AzureKeyCredential
from azure.search.documents.indexes import SearchIndexClient
from docx import Document as DocxDocument
from llama_index.core import Settings, StorageContext, VectorStoreIndex
from llama_index.core.schema import Document
from llama_index.embeddings.azure_openai import AzureOpenAIEmbedding
from llama_index.vector_stores.azureaisearch import (
    AzureAISearchVectorStore,
    IndexManagement,
)
from app_logger import setup_logger

logger, _ = setup_logger(name="llama-indexer")

CHUNK_SIZE = int(os.getenv("INDEX_CHUNK_SIZE", "1200"))
CHUNK_OVERLAP = int(os.getenv("INDEX_CHUNK_OVERLAP", "200"))
EMBED_BATCH_SIZE = int(os.getenv("EMBED_BATCH_SIZE", "16"))
INDEX_VERSION = os.getenv("INDEX_VERSION", "v2")

DEFAULT_EMBEDDING_MODEL = os.getenv(
    "AZURE_OPENAI_EMBED_MODEL",
    "text-embedding-3-large",
)
DEFAULT_EMBEDDING_DIMENSION = int(
    os.getenv("AZURE_OPENAI_EMBED_DIMENSION", "3072")
)
DEFAULT_INDEX_METADATA_PATH = os.getenv(
    "INDEX_METADATA_PATH",
    ".index_metadata.json",
)

SUPPORTED_EXTENSIONS = frozenset({".pdf", ".docx", ".txt", ".md", ".csv"})


@dataclass(frozen=True)
class IndexRuntime:
    """Compatibility runtime returned by ``init_embedding_and_vectorstore``."""

    embedding: Any
    vector_store: AzureAISearchVectorStore


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def now_iso() -> str:
    """Return a timezone-aware UTC timestamp."""
    return _utc_now().isoformat().replace("+00:00", "Z")


def compute_checksum_bytes(data: bytes) -> str:
    """Return a SHA-256 checksum for bytes."""
    return hashlib.sha256(data).hexdigest()


def compute_checksum_file(path: str) -> str:
    """Compute SHA-256 without loading the complete file into memory."""
    digest = hashlib.sha256()
    with open(path, "rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def make_doc_id(source: str) -> str:
    """Create a deterministic identifier from a stable source identifier."""
    return str(uuid.uuid5(uuid.NAMESPACE_URL, source))


def _validate_chunking(chunk_size: int, overlap: int) -> None:
    if isinstance(chunk_size, bool) or not isinstance(chunk_size, int):
        raise ValueError("chunk_size must be an integer.")
    if isinstance(overlap, bool) or not isinstance(overlap, int):
        raise ValueError("overlap must be an integer.")
    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than zero.")
    if overlap < 0:
        raise ValueError("overlap must be non-negative.")
    if overlap >= chunk_size:
        raise ValueError("overlap must be smaller than chunk_size.")


def chunk_text_simple(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> List[Tuple[str, int, int]]:
    """Chunk text using deterministic character windows with overlap."""
    _validate_chunking(chunk_size, overlap)

    if not text:
        return []

    chunks: List[Tuple[str, int, int]] = []
    step = chunk_size - overlap

    for start in range(0, len(text), step):
        end = min(start + chunk_size, len(text))
        chunks.append((text[start:end], start, end))
        if end >= len(text):
            break

    return chunks


def extract_text_from_pdf(path: str) -> Tuple[str, int]:
    """Extract text from a PDF and always close the PyMuPDF document."""
    document = fitz.open(path)
    try:
        parts = [page.get_text("text") for page in document]
        return "\n".join(parts), len(document)
    finally:
        document.close()


def extract_text_from_docx(path: str) -> Tuple[str, None]:
    document = DocxDocument(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    return "\n".join(paragraphs), None


def extract_text_from_txt(path: str) -> Tuple[str, None]:
    with open(path, "r", encoding="utf-8", errors="replace") as handle:
        return handle.read(), None


def extract_text_from_csv(
    path: str,
    text_columns: Optional[List[str]] = None,
    max_rows: Optional[int] = None,
) -> Tuple[str, None]:
    dataframe = pd.read_csv(path)
    return extract_text_from_dataframe(
        dataframe,
        text_columns=text_columns,
        max_rows=max_rows,
    )


def extract_text_from_dataframe(
    dataframe: pd.DataFrame,
    text_columns: Optional[List[str]] = None,
    max_rows: Optional[int] = None,
) -> Tuple[str, None]:
    if not isinstance(dataframe, pd.DataFrame):
        raise TypeError("dataframe must be a pandas DataFrame.")

    selected = dataframe
    if text_columns:
        missing = [column for column in text_columns if column not in selected.columns]
        if missing:
            raise ValueError(f"Unknown DataFrame columns: {missing}")
        selected = selected.loc[:, text_columns]

    if max_rows is not None:
        if isinstance(max_rows, bool) or not isinstance(max_rows, int) or max_rows < 0:
            raise ValueError("max_rows must be a non-negative integer.")
        selected = selected.head(max_rows)

    return selected.to_csv(index=False), None


def _guess_mime(path: str) -> str:
    extension = Path(path).suffix.lower()
    explicit = {
        ".pdf": "application/pdf",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".csv": "text/csv",
        ".docx": (
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    }
    return explicit.get(extension) or mimetypes.guess_type(path)[0] or (
        "application/octet-stream"
    )


def build_doc_metadata_for_file(
    path: str,
    checksum: str,
    extra: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    file_path = Path(path)
    stat = file_path.stat()

    metadata: Dict[str, Any] = {
        "doc_id": make_doc_id(str(file_path.resolve())),
        "source_path": str(file_path.resolve()),
        "filename": file_path.name,
        "checksum": checksum,
        "uploaded_date": now_iso(),
        "last_modified": datetime.fromtimestamp(
            stat.st_mtime,
            tz=timezone.utc,
        ).isoformat().replace("+00:00", "Z"),
        "file_size": stat.st_size,
        "mime_type": _guess_mime(path),
        "index_version": INDEX_VERSION,
        "ingest_pipeline": "llama-indexer-v2",
    }

    if extra:
        metadata.update(extra)

    return metadata


def create_documents_from_text(
    text: str,
    base_meta: Dict[str, Any],
    *,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP,
) -> List[Document]:
    if not isinstance(text, str):
        raise TypeError("text must be a string.")
    if "doc_id" not in base_meta:
        raise ValueError("base_meta must contain doc_id.")

    documents: List[Document] = []

    for index, (chunk, start, end) in enumerate(
        chunk_text_simple(text, chunk_size=chunk_size, overlap=overlap)
    ):
        chunk_id = f"{base_meta['doc_id']}::chunk::{index}"
        metadata = dict(base_meta)
        metadata.update(
            {
                "chunk_id": chunk_id,
                "chunk_start_offset": start,
                "chunk_end_offset": end,
                "chunk_start_page": base_meta.get("chunk_start_page"),
                "chunk_end_page": base_meta.get("chunk_end_page"),
                "indexed_at": now_iso(),
            }
        )
        documents.append(
            Document(
                text=chunk,
                metadata=metadata,
                id_=chunk_id,
            )
        )

    return documents


def _require_env(name: str) -> str:
    value = os.getenv(name)
    if not value or not value.strip():
        raise ValueError(f"Required environment variable '{name}' is missing.")
    return value.strip()


def init_embedding_and_vectorstore() -> Tuple[Any, AzureAISearchVectorStore, None]:
    """Create the embedding model and Azure AI Search vector store.

    The third return value is retained as ``None`` for compatibility with the
    old ``(embedding, vector_store, service_context)`` tuple. Current
    LlamaIndex uses ``Settings`` instead of ``ServiceContext``.
    """
    deployment = os.getenv(
        "AZURE_OPENAI_EMBED_DEPLOYMENT",
        DEFAULT_EMBEDDING_MODEL,
    )
    endpoint = _require_env("AZURE_OPENAI_ENDPOINT")
    api_key = os.getenv("AZURE_OPENAI_API_KEY")

    embedding_kwargs: Dict[str, Any] = {
        "model": DEFAULT_EMBEDDING_MODEL,
        "deployment_name": deployment,
        "azure_endpoint": endpoint,
        "api_version": _require_env("AZURE_OPENAI_API_VERSION"),
        "embed_batch_size": EMBED_BATCH_SIZE,
    }
    if api_key:
        embedding_kwargs["api_key"] = api_key
    else:
        # AzureOpenAIEmbedding can use an Azure AD token provider when supplied
        # by the surrounding application. This standalone module preserves the
        # environment-key path and fails clearly when no key is available.
        raise ValueError(
            "AZURE_OPENAI_API_KEY is required by this standalone indexer. "
            "Use the application's credential-aware embedding loader when "
            "managed identity is required."
        )

    embedding = AzureOpenAIEmbedding(**embedding_kwargs)
    Settings.embed_model = embedding
    Settings.chunk_size = CHUNK_SIZE
    Settings.chunk_overlap = CHUNK_OVERLAP

    service_endpoint = os.getenv("AZURE_SEARCH_ENDPOINT")
    if not service_endpoint:
        service_name = _require_env("AZURE_SEARCH_SERVICE_NAME")
        service_endpoint = f"https://{service_name}.search.windows.net"

    search_key = _require_env("AZURE_SEARCH_API_KEY")
    index_name = _require_env("AZURE_SEARCH_INDEX_NAME")

    client = SearchIndexClient(
        endpoint=service_endpoint,
        credential=AzureKeyCredential(search_key),
    )

    vector_store = AzureAISearchVectorStore(
        search_or_index_client=client,
        index_name=index_name,
        index_management=IndexManagement.CREATE_IF_NOT_EXISTS,
        id_field_key="id",
        chunk_field_key="chunk",
        embedding_field_key="embedding",
        embedding_dimensionality=DEFAULT_EMBEDDING_DIMENSION,
        metadata_string_field_key="metadata",
        doc_id_field_key="doc_id",
        vector_algorithm_type=os.getenv(
            "AZURE_SEARCH_VECTOR_ALGORITHM",
            "exhaustiveKnn",
        ),
    )

    return embedding, vector_store, None


def _load_metadata(path: str) -> Dict[str, Any]:
    metadata_path = Path(path)
    if not metadata_path.exists():
        return {}

    try:
        with metadata_path.open("r", encoding="utf-8") as handle:
            value = json.load(handle)
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"Invalid index metadata file: {metadata_path}") from exc

    if not isinstance(value, dict):
        raise ValueError(f"Index metadata must be a JSON object: {metadata_path}")

    return value


def _save_metadata(path: str, metadata: Dict[str, Any]) -> None:
    metadata_path = Path(path)
    metadata_path.parent.mkdir(parents=True, exist_ok=True)

    temporary = metadata_path.with_suffix(metadata_path.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8") as handle:
        json.dump(metadata, handle, indent=2, sort_keys=True)
        handle.write("\n")
    temporary.replace(metadata_path)


def _is_unchanged(
    metadata: Dict[str, Any],
    doc_id: str,
    checksum: str,
) -> bool:
    record = metadata.get(doc_id)
    return (
        isinstance(record, dict)
        and record.get("checksum") == checksum
        and record.get("index_version") == INDEX_VERSION
    )


def _resolve_runtime(
    vector_store: Any = None,
    service_context: Any = None,
) -> Tuple[Any, Any]:
    del service_context  # Compatibility-only argument; ServiceContext is removed.
    if vector_store is not None:
        return vector_store, Settings.embed_model

    embedding, store, _ = init_embedding_and_vectorstore()
    return store, embedding


def _delete_existing_document(vector_store: Any, doc_id: str) -> None:
    delete_method = getattr(vector_store, "delete", None)
    if not callable(delete_method):
        return

    try:
        delete_method(doc_id)
    except TypeError:
        # Some vector-store versions use keyword naming.
        try:
            delete_method(ref_doc_id=doc_id)
        except Exception:
            logger.warning(
                "Existing document cleanup was unavailable for doc_id=%s",
                doc_id,
                exc_info=True,
            )
    except Exception:
        logger.warning(
            "Existing document cleanup failed for doc_id=%s",
            doc_id,
            exc_info=True,
        )


def _index_documents(
    documents: Sequence[Document],
    vector_store: Any,
) -> None:
    if not documents:
        return

    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    VectorStoreIndex(
        nodes=list(documents),
        storage_context=storage_context,
        embed_model=Settings.embed_model,
        insert_batch_size=EMBED_BATCH_SIZE,
        show_progress=False,
    )


def index_file(
    path: str,
    force_reindex: bool = False,
    vector_store: Any = None,
    service_context: Any = None,
) -> Dict[str, Any]:
    """Extract, chunk, embed, and upsert one supported file."""
    file_path = Path(path)

    if not file_path.is_file():
        raise FileNotFoundError(f"File not found: {path}")

    extension = file_path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {extension or '<none>'}. "
            f"Supported types: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    checksum = compute_checksum_file(str(file_path))
    metadata = build_doc_metadata_for_file(str(file_path), checksum)

    metadata_path = os.getenv(
        "INDEX_METADATA_PATH",
        DEFAULT_INDEX_METADATA_PATH,
    )
    index_metadata = _load_metadata(metadata_path)

    if not force_reindex and _is_unchanged(
        index_metadata,
        metadata["doc_id"],
        checksum,
    ):
        return {
            "doc_id": metadata["doc_id"],
            "chunks_indexed": 0,
            "checksum": checksum,
            "status": "skipped",
            "reason": "unchanged",
        }

    if extension == ".pdf":
        text, page_count = extract_text_from_pdf(str(file_path))
        metadata["page_count"] = page_count
    elif extension == ".docx":
        text, _ = extract_text_from_docx(str(file_path))
    elif extension in {".txt", ".md"}:
        text, _ = extract_text_from_txt(str(file_path))
    else:
        text, _ = extract_text_from_csv(str(file_path))

    if not text.strip():
        return {
            "doc_id": metadata["doc_id"],
            "chunks_indexed": 0,
            "checksum": checksum,
            "status": "skipped",
            "reason": "empty_document",
        }

    documents = create_documents_from_text(text, metadata)
    store, _ = _resolve_runtime(vector_store, service_context)

    if force_reindex or metadata["doc_id"] in index_metadata:
        _delete_existing_document(store, metadata["doc_id"])

    logger.info(
        "Indexing doc_id=%s filename=%s chunks=%d",
        metadata["doc_id"],
        metadata["filename"],
        len(documents),
    )

    _index_documents(documents, store)

    index_metadata[metadata["doc_id"]] = {
        "checksum": checksum,
        "index_version": INDEX_VERSION,
        "source_path": metadata["source_path"],
        "filename": metadata["filename"],
        "chunks_indexed": len(documents),
        "indexed_at": now_iso(),
    }
    _save_metadata(metadata_path, index_metadata)

    return {
        "doc_id": metadata["doc_id"],
        "chunks_indexed": len(documents),
        "checksum": checksum,
        "status": "indexed",
    }


def index_dataframe(
    dataframe: pd.DataFrame,
    name: str = "dataframe",
    text_columns: Optional[List[str]] = None,
    vector_store: Any = None,
    service_context: Any = None,
) -> Dict[str, Any]:
    """Index a DataFrame using the same deterministic document pipeline."""
    if not isinstance(name, str) or not name.strip():
        raise ValueError("name must be a non-empty string.")

    text, _ = extract_text_from_dataframe(
        dataframe,
        text_columns=text_columns,
    )
    checksum = compute_checksum_bytes(text.encode("utf-8"))

    metadata = {
        "doc_id": make_doc_id(f"{name}:{checksum}"),
        "source_path": name,
        "filename": name,
        "checksum": checksum,
        "uploaded_date": now_iso(),
        "last_modified": now_iso(),
        "file_size": len(text.encode("utf-8")),
        "mime_type": "text/csv",
        "index_version": INDEX_VERSION,
        "ingest_pipeline": "dataframe-csv",
    }

    documents = create_documents_from_text(text, metadata)
    store, _ = _resolve_runtime(vector_store, service_context)

    _index_documents(documents, store)

    return {
        "doc_id": metadata["doc_id"],
        "chunks_indexed": len(documents),
        "checksum": checksum,
        "status": "indexed",
    }


def semantic_search(
    query: str,
    top_k: int = 5,
    vector_store: Any = None,
    service_context: Any = None,
):
    """Query an existing vector store with the requested top-k."""
    if not isinstance(query, str) or not query.strip():
        raise ValueError("query must be a non-empty string.")
    if isinstance(top_k, bool) or not isinstance(top_k, int) or top_k <= 0:
        raise ValueError("top_k must be a positive integer.")

    store, _ = _resolve_runtime(vector_store, service_context)

    index = VectorStoreIndex.from_vector_store(
        vector_store=store,
        embed_model=Settings.embed_model,
    )
    return index.as_query_engine(similarity_top_k=top_k).query(query)


def index_path(
    path: str,
    recursive: bool = True,
    force: bool = False,
) -> List[Dict[str, Any]]:
    """Index one file or all supported files under a directory."""
    target = Path(path)
    if not target.exists():
        raise FileNotFoundError(f"Path not found: {path}")

    _, vector_store, _ = init_embedding_and_vectorstore()
    files = (
        [target]
        if target.is_file()
        else (
            target.rglob("*")
            if recursive
            else target.glob("*")
        )
    )

    results: List[Dict[str, Any]] = []
    for file_path in sorted(files):
        if not file_path.is_file() or file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue

        try:
            results.append(
                index_file(
                    str(file_path),
                    force_reindex=force,
                    vector_store=vector_store,
                )
            )
        except Exception:
            logger.exception("Failed to index file: %s", file_path)
            results.append(
                {
                    "source_path": str(file_path),
                    "status": "failed",
                }
            )

    return results
