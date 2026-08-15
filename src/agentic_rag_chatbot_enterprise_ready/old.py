
# # # # import os
# # # # from pathlib import Path
# # # # from pdfplumber import open as pdfplumber_open
# # # # from PIL import Image
# # # # import pytesseract
# # # # import pandas as pd
# # # # from typing import Dict, Any, List

# # # # # --- STAGE 1: Core Utility Functions ---

# # # # def detect_file_type(file_path: Path) -> str:
# # # #     """
# # # #     Detects the file format and type (PDF, image, structured, etc.)
# # # #     Using file signature checks (python-magic is better for this, 
# # # #     but we'll use basic extensions for simplicity).
# # # #     """
# # # #     ext = file_path.suffix.lower()
    
# # # #     if ext == '.pdf':
# # # #         return "PDF"
# # # #     elif ext in ['.jpg', '.jpeg', '.png']:
# # # #         return "IMAGE"
# # # #     elif ext in ['.xlsx', '.xls']:
# # # #         return "STRUCTURED_XL"
# # # #     elif ext in ['.csv', '.txt']:
# # # #         return "UNSTRUCTURED_TEXT"
# # # #     else:
# # # #         return "UNKNOWN"

# # # # def preprocess_document(file_path: Path) -> str:
# # # #     """
# # # #     Initial loading based on detected type. Returns raw content or path for processing.
# # # #     """
# # # #     doc_type = detect_file_type(file_path)
# # # #     print(f"-> Detected Type: {doc_type}")

# # # #     if doc_type == "PDF":
# # # #         return process_pdf(file_path)
# # # #     elif doc_type == "IMAGE":
# # # #         return process_image(file_path)
# # # #     elif doc_type == "STRUCTURED_XL":
# # # #         return process_structured_excel(file_path)
# # # #     elif doc_type == "UNSTRUCTURED_TEXT":
# # # #         return file_path.read_text()
# # # #     else:
# # # #         raise ValueError(f"Unsupported document type: {doc_type}")

# # # # # --- STAGE 2: Format-Specific Processors ---

# # # # def process_pdf(file_path: Path) -> str:
# # # #     """Handles PDF files (text extraction and structural analysis)."""
# # # #     text_content = []
# # # #     try:
# # # #         with pdfplumber_open(file_path) as pdf:
# # # #             for page in pdf.pages:
# # # #                 text = page.extract_text()
# # # #                 if text:
# # # #                     text_content.append(text)
        
# # # #         # For advanced use: Use pdfplumber to extract tables specifically
# # # #         # tables = []
# # # #         # for page in pdf.pages:
# # # #         #     tables.extend(page.extract_tables())
            
# # # #         return "\n".join(text_content)
# # # #     except Exception as e:
# # # #         print(f"Warning: PDF processing failed, attempting OCR fallback. Error: {e}")
# # # #         # FALLBACK: If simple extraction fails, run OCR on the PDF pages
# # # #         return "OCR_FALLBACK_NEEDED" # Trigger dedicated OCR step

# # # # def process_image(file_path: Path) -> str:
# # # #     """Handles image files using OCR (Tesseract or Cloud API)."""
# # # #     print("   [OCR] Running OCR on image...")
# # # #     try:
# # # #         # Use PIL to open the image
# # # #         img = Image.open(file_path)
        
# # # #         # Local Tesseract OCR implementation (simple, less accurate than cloud services)
# # # #         raw_text = pytesseract.image_to_string(img)
# # # #         return raw_text
# # # #     except Exception as e:
# # # #         print(f"Error during local OCR: {e}")
# # # #         # *** PRODUCTION NOTE: Replace this block with AWS Textract or Google Vision API calls ***
# # # #         return f"[OCR FAILED - Use Cloud API]: {file_path.name}"

# # # # def process_structured_excel(file_path: Path) -> str:
# # # #     """Handles structured data (Excel/CSV) and converts it to structured text/markdown."""
# # # #     print("   [Extraction] Reading structured data...")
# # # #     try:
# # # #         if file_path.suffix.lower() == '.csv':
# # # #              df = pd.read_csv(file_path)
# # # #         else:
# # # #              df = pd.read_excel(file_path)

# # # #         # Convert DataFrame to a readable markdown/text format
# # # #         return "--- STRUCTURED DATA (FROM EXCEL/CSV) ---\n" + df.to_markdown(index=False)
# # # #     except Exception as e:
# # # #         return f"ERROR PROCESSING STRUCTURED DATA: {e}"

# # # # # --- STAGE 3: Advanced Parsing and Chunking ---

# # # # def parse_and_extract_context(raw_text: str) -> List[Dict[str, Any]]:
# # # #     """
# # # #     Takes raw text and uses NLP/LLMs to extract meaning, chunks, and metadata.
# # # #     This is the "intelligence" layer.
# # # #     """
# # # #     print("   [Parsing] Running NLP extraction and chunking...")
    
# # # #     # 1. Chunking: Break the massive text block into smaller, digestible chunks (e.g., 500 tokens)
# # # #     # This is crucial for effective vector storage.
# # # #     # For demonstration, we use simple line splits.
# # # #     chunks = [chunk.strip() for chunk in raw_text.split('\n\n') if chunk.strip()]
    
# # # #     processed_chunks = []
    
# # # #     # 2. Intelligent Extraction (Placeholder for LLM calls)
# # # #     for i, chunk in enumerate(chunks):
# # # #         # In a real system, you would pass this chunk to an LLM (e.g., GPT-4)
# # # #         # with a detailed prompt:
# # # #         # PROMPT = "Analyze the following text chunk. Identify the document type, date, 
# # # #         # customer name, and invoice ID. Output a JSON object."
# # # #         # metadata = call_llm_api(chunk, prompt=PROMPT)
        
# # # #         # Simplified simulated extraction:
# # # #         metadata = {
# # # #             "source_chunk_id": f"chunk_{i}",
# # # #             "document_type_guess": "Invoice" if "invoice" in chunk.lower() else "General Document",
# # # #             "page_context": "Paging done by parser" # Would come from the PDF parser
# # # #         }
        
# # # #         processed_chunks.append({
# # # #             "chunk_text": chunk,
# # # #             "metadata": metadata
# # # #         })
        
# # # #     return processed_chunks

# # # # def embed_chunk(chunk: str) -> str:
# # # #     """
# # # #     Generates a vector embedding (the numerical representation of the text meaning).
# # # #     """
# # # #     # *** PRODUCTION NOTE: Replace with actual OpenAI/Cohere/HuggingFace embedding call ***
# # # #     # Example: embedding_model.get_embedding(chunk)
# # # #     return f"vector_embedding_for_{len(chunk)}chars"

# # # # # --- STAGE 4: Storage and Indexing ---

# # # # def store_data(processed_chunks: List[Dict[str, Any]], source_file: Path):
# # # #     """
# # # #     Stores the data: vectors in Vector DB, and full data in traditional DB.
# # # #     """
# # # #     print("\n--- STORAGE & INDEXING ---")
    
# # # #     # 1. Vector Database Storage (Semantic Search)
# # # #     print("   [Vector DB] Indexing chunks...")
# # # #     # Imagine calling Pinecone/Chroma SDK here
# # # #     # for chunk in processed_chunks:
# # # #     #     vector = embed_chunk(chunk['chunk_text'])
# # # #     #     db_client.upsert(vector=vector, metadata=chunk['metadata'], id=...)
    
# # # #     # 2. Non-Vector Database Storage (Record Keeping/Filtering)
# # # #     print("   [Non-Vector DB] Storing metadata...")
# # # #     # In a real scenario, you might store the structured metadata 
# # # #     # (Customer ID, Date, etc.) in PostgreSQL or MongoDB for fast filtering.
    
# # # #     # Example of simulating the DB record generation
# # # #     db_records = []
# # # #     for chunk in processed_chunks:
# # # #         record = {
# # # #             "source_file": source_file.name,
# # # #             "chunk_text": chunk['chunk_text'],
# # # #             "document_type": chunk['metadata']['document_type_guess'],
# # # #             "vector_id": f"sim_vec_{hash(chunk['chunk_text']) % 100}"
# # # #         }
# # # #         db_records.append(record)
        
# # # #     print(f"   [SUCCESS] Indexed {len(db_records)} chunks and saved metadata.")
# # # #     # (Here you would use SQLAlchemy or similar ORM to connect and save these records)


# # # # # --- THE MAIN ORCHESTRATOR ---

# # # # def document_digitization_pipeline(file_path: Path):
# # # #     """
# # # #     The main pipeline function to ingest, process, parse, and store a document.
# # # #     """
# # # #     print("=============================================")
# # # #     print(f"🚀 Starting Digitization Pipeline for: {file_path.name}")
# # # #     print("=============================================")
    
# # # #     try:
# # # #         # STAGE 1 & 2: Preprocessing and Content Extraction
# # # #         raw_text_content = preprocess_document(file_path)
        
# # # #         if raw_text_content.startswith("[OCR FAILED"):
# # # #              print("❌ Pipeline halted due to critical OCR failure.")
# # # #              return

# # # #         # STAGE 3: Parsing, Chunking, and Embedding
# # # #         processed_chunks = parse_and_extract_context(raw_text_content)
        
# # # #         if not processed_chunks:
# # # #             print("⚠️ Warning: No usable content was extracted.")
# # # #             return
            
# # # #         # STAGE 4: Storage and Indexing
# # # #         store_data(processed_chunks, file_path)
        
# # # #     except Exception as e:
# # # #         print(f"\n🚨 FATAL PIPELINE ERROR: {e}")

# # # # # --- EXAMPLE USAGE ---

# # # # if __name__ == "__main__":
# # # #     # --- SET UP DUMMY FILES FOR TESTING ---
# # # #     # In a real scenario, you would load files from a directory list.
    
# # # #     # 1. Dummy Text File (Unstructured)
# # # #     Path("test_doc_1.txt").write_text("""
# # # #     Client Invoice Report - 2023-10-27
# # # #     Dear John Doe,
# # # #     Please find attached the invoice details for Q4 services. 
# # # #     Total due amount is $1,250.
# # # #     Reference ID: INV-447B.
# # # #     """)
    
# # # #     # 2. Dummy Image File (Requires a PNG/JPG image file)
# # # #     # NOTE: For real testing, ensure you have a test image named 'test_image.jpg' 
# # # #     # in the same directory.
# # # #     image_path = Path("test_image.jpg") 
# # # #     if not image_path.exists():
# # # #          print("\n[SKIP]: Please place a test image (e.g., receipt.png) named 'test_image.jpg' to run the image test.")
# # # #     else:
# # # #         document_digitization_pipeline(image_path)

# # # #     # 3. Dummy PDF File (Requires a PDF file named 'test_doc_2.pdf')
# # # #     # NOTE: For real testing, ensure you have a test PDF named 'test_doc_2.pdf'
# # # #     pdf_path = Path("test_doc_2.pdf")
# # # #     if pdf_path.exists():
# # # #         document_digitization_pipeline(pdf_path)
        
# # # #     # 4. Dummy Structured File (CSV/Excel)
# # # #     dummy_df = pd.DataFrame({'Product': ['A', 'B'], 'Quantity': [10, 5], 'Price': [2.5, 10.0]})
# # # #     Path("test_doc_3.csv").write_csv(dummy_df, index=False)
# # # #     document_digitization_pipeline(Path("test_doc_3.csv"))

# # # import os
# # # import logging
# # # from datetime import datetime
# # # import numpy as np
# # # import pandas as pd
# # # from docdetector import DocumentDetector
# # # import cv2
# # # from OCR import OCR  # Using Tesseract with Python
# # # from PyPDF2 import PdfWriter
# # # from FAISS import Faiss
# # # from pymongo import MongoClient


# # # logging.basicConfig(level=logging.INFO)
# # # logger = logging.getLogger(__name__)

# # # def convert_document(path, output_format='image'):
# # #     """Converts a document to the specified format."""
# # #     if output_format == 'image':
# # #         # Convert to grayscale
# # #         img = cv2.imread(path)
# # #         img_gray = cv2.cvtColor(img, cv2.COLOR_RGB_to_GRAY)
# # #         cv2.imwrite('document_image.png', img_gray)
# # #         return 'document_image.png'
# # #     elif output_format == 'pdf':
# # #         # Convert using PyPDF2
# # #         writer = PdfWriter()
# # #         writer.write('document.pdf', 'document_content')
# # #         return 'document.pdf'
# # #     elif output_format == 'ocr':
# # #         # Convert to OCR text
# # #         if not os.path.exists('tesseract.exe'):
# # #             logger.error("Tesseract not found. Please install it.")
# # #             raise FileNotFoundError("Tesseract not found")
# # #         ocr = OCR()
# # #         text = ocr.detect_and_extract(path)
# # #         return text
# # #     else:
# # #         raise ValueError("Invalid output format")

# # # def extract_structured_data(document_path, detector, extractor):
# # #     """Extracts structured data from a document."""
# # #     doc_type, _ = detect_document_type(document_path)
# # #     if doc_type == 'ocr':
# # #         text = extractor.process(document_path)
# # #         logger.info(f"OCR text extracted: {text}")
# # #         return text
# # #     else:
# # #         data = extractor.extract_data(document_path, doc_type)
# # #         logger.debug(f"Structured data extracted: {data}")
# # #         return data

# # # def main():
# # #     # Configuration
# # #     config = {
# # #         'input_folder': 'documents',
# # #         'output_folder': 'processed_documents',
# # #         'database': {
# # #             'uri': 'sqlite:///database.db',
# # #             'collection': 'documents'
# # #         }
# # #     }

# # #     # Initialize database
# # #     db = Database(config['database'])

# # #     # Detector and extractor setup
# # #     detector = DocumentDetector()
# # #     extractor = DocumentExtractor()

# # #     # Initialize logger
# # #     logger.info("Starting document digitization process.")

# # #     # Process documents
# # #     for filename in os.listdir(config['input_folder']):
# # #         if not filename.endswith(('.pdf', '.docx', '.xlsx', '.txt')):
# # #             continue
# # #         full_path = os.path.join(config['input_folder'], filename)
# # #         logger.info(f"Processing: {filename}")
# # #         try:
# # #             # Detect document type
# # #             doc_type = detector.detect(full_path)
            
# # #             # Convert document
# # #             converted_path = convert_document(full_path)
            
# # #             # Extract structured data
# # #             data = extract_structured_data(converted_path, detector, extractor)
            
# # #             # Store data
# # #             store_data(data, db)
            
# # #             logger.info(f"Document {filename} processed successfully.")
# # #         except Exception as e:
# # #             logger.error(f"Error processing {filename}: {str(e)}")
# # #             logger.error(f"Stack trace:\n{traceback.format_exc()}")

# # # if __name__ == "__main__":
# # #     main()

# # # import os
# # # import logging
# # # from datetime import datetime
# # # import numpy as np
# # # import pandas as pd
# # # from docdetector import DocumentDetector
# # # import cv2
# # # from OCR import OCR  # Using Tesseract with Python
# # # from PyPDF2 import PdfWriter
# # # from FAISS import Faiss
# # # from pymongo import MongoClient

# # # # Configure logging
# # # logging.basicConfig(level=logging.INFO)
# # # logger = logging.getLogger(__name__)

# # # class DocumentDigitizationSystem:
# # #     def __init__(self, config):
# # #         self.config = config
# # #         self.database = Database(config['database'])
# # #         self.detector = DocumentDetector()
# # #         self.extractor = DocumentExtractor()
        
# # #     def detect_document_type(self, path):
# # #         """Detects the type and format of a document."""
# # #         return self.detector.detect(path)
    
# # #     def convert_document(self, path, output_format='image'):
# # #         """Converts a document to the specified format."""
# # #         if output_format == 'image':
# # #             # Convert to grayscale
# # #             img = cv2.imread(path)
# # #             img_gray = cv2.cvtColor(img, cv2.COLOR_RGB_to_GRAY)
# # #             cv2.imwrite('document_image.png', img_gray)
# # #             return 'document_image.png'
# # #         elif output_format == 'pdf':
# # #             # Convert using PyPDF2
# # #             writer = PdfWriter()
# # #             writer.write('document.pdf', 'document_content')
# # #             return 'document.pdf'
# # #         elif output_format == 'ocr':
# # #             # Convert to OCR text using Tesseract
# # #             if not os.path.exists('tesseract.exe'):
# # #                 logger.error("Tesseract not found. Please install it.")
# # #                 raise FileNotFoundError("Tesseract not found")
# # #             ocr = OCR()
# # #             text = ocr.detect_and_extract(path)
# # #             return text
# # #         else:
# # #             raise ValueError("Invalid output format")
    
# # #     def extract_structured_data(self, document_path, doc_type):
# # #         """Extracts structured data from a document."""
# # #         if doc_type == 'ocr':
# # #             text = self.extractor.process(document_path)
# # #             logger.info(f"OCR text extracted: {text}")
# # #             return text
# # #         else:
# # #             data = self.extractor.extract_data(document_path, doc_type)
# # #             logger.debug(f"Structured data extracted: {data}")
# # #             return data
    
# # #     def store_data(self, data, collection='documents'):
# # #         """Stores data in the database."""
# # #         self.database.store(data, collection)
    
# # #     def query_data(self, query, collection='documents'):
# # #         """Queries the database."""
# # #         return self.database.query(query, collection)

# # # class Database:
# # #     def __init__(self, config):
# # #         self.config = config
# # #         self._connect_to_database()
    
# # #     def _connect_to_database(self):
# # #         """Connects to the database."""
# # #         if self.config['database']['uri'].startswith('sqlite'):
# # #             self._connect_sqlite()
# # #         else:
# # #             self._connect_mongodb()
    
# # #     def _connect_sqlite(self):
# # #         from sqlite3 import connect
# # #         self.conn = connect(self.config['database']['uri'])
# # #         self.conn.execute("CREATE TABLE IF NOT EXISTS documents (id INTEGER PRIMARY KEY, data TEXT, processed_at DATETIME DEFAULT CURRENT_TIMESTAMP)")
    
# # #     def _connect_mongodb(self):
# # #         self.client = MongoClient(self.config['database']['uri'])
# # #         self.db = self.client[self.config['database']['collection']]
# # #         self.db.create_index("text", "data", unique=False)
    
# # #     def store(self, data, collection='documents'):
# # #         """Stores data in the database."""
# # #         document_id = str(uuid.uuid4())
# # #         self.db[collection].insert_one({
# # #             "_id": document_id,
# # #             "data": data,
# # #             "processed_at": datetime.now()
# # #         })
    
# # #     def query(self, query, collection='documents'):
# # #         """Queries the database."""
# # #         return self.db[collection].find({"$text": {"$search": query}})
    
# # #     def close(self):
# # #         """Closes the database connection."""
# # #         if self.config['database']['uri'].startswith('sqlite'):
# # #             self.conn.close()
# # #         else:
# # #             self.client.close()

# # # def main():
# # #     # Configuration
# # #     config = {
# # #         'input_folder': 'documents',
# # #         'output_folder': 'processed_documents',
# # #         'database': {
# # #             'uri': 'sqlite:///database.db',
# # #             'collection': 'documents'
# # #         }
# # #     }

# # #     # Initialize database
# # #     db = Database(config['database'])

# # #     # Detector and extractor setup
# # #     detector = DocumentDetector()
# # #     extractor = DocumentExtractor()

# # #     # Process documents
# # #     for filename in os.listdir(config['input_folder']):
# # #         if not filename.endswith(('.pdf', '.docx', '.xlsx', '.txt')):
# # #             continue
# # #         full_path = os.path.join(config['input_folder'], filename)
# # #         logger.info(f"Processing: {filename}")
# # #         try:
# # #             # Detect document type
# # #             doc_type = detector.detect(full_path)
            
# # #             # Convert document
# # #             converted_path = db.convert_document(full_path)
            
# # #             # Extract structured data
# # #             data = db.extract_structured_data(converted_path, doc_type)
            
# # #             # Store data
# # #             db.store_data(data)
            
# # #             logger.info(f"Document {filename} processed successfully.")
# # #         except Exception as e:
# # #             logger.error(f"Error processing {filename}: {str(e)}")
# # #             logger.error(f"Stack trace:\n{traceback.format_exc()}")

# # # if __name__ == "__main__":
# # #     main()
# # import os
# # import logging
# # from pathlib import Path
# # from typing import Dict, Any

# # import pandas as pd
# # import fitz  # PyMuPDF
# # import docx

# # # Configure logging
# # logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
# # logger = logging.getLogger(__name__)

# # class DocumentDigitizer:
# #     """
# #     Handles extraction of text and metadata from various document formats 
# #     (PDF, DOCX, CSV/Excel, Text) into a unified digitized format.
# #     """
# #     def __init__(self):
# #         logger.info("Initialized DocumentDigitizer.")

# #     def detect_file_type(self, file_path: Path) -> str:
# #         """Detects the file format based on extension."""
# #         ext = file_path.suffix.lower()
# #         if ext == '.pdf':
# #             return "PDF"
# #         elif ext in ['.docx']:
# #             return "DOCX"
# #         elif ext in ['.xlsx', '.xls', '.csv']:
# #             return "STRUCTURED"
# #         elif ext in ['.txt', '.md']:
# #             return "TEXT"
# #         else:
# #             return "UNKNOWN"

# #     def process_pdf(self, file_path: Path) -> str:
# #         """Extracts text from PDF files using PyMuPDF (fitz)."""
# #         text_content = []
# #         try:
# #             with fitz.open(file_path) as doc:
# #                 for page in doc:
# #                     text_content.append(page.get_text())
# #             return "\n\n".join(text_content)
# #         except Exception as e:
# #             logger.error(f"Error processing PDF {file_path}: {e}")
# #             return ""

# #     def process_docx(self, file_path: Path) -> str:
# #         """Extracts text from DOCX files using python-docx."""
# #         text_content = []
# #         try:
# #             doc = docx.Document(file_path)
# #             for para in doc.paragraphs:
# #                 if para.text.strip():
# #                     text_content.append(para.text)
# #             return "\n\n".join(text_content)
# #         except Exception as e:
# #             logger.error(f"Error processing DOCX {file_path}: {e}")
# #             return ""

# #     def process_structured(self, file_path: Path) -> str:
# #         """Extracts text from CSV/Excel files using pandas."""
# #         try:
# #             if file_path.suffix.lower() == '.csv':
# #                 df = pd.read_csv(file_path)
# #             else:
# #                 df = pd.read_excel(file_path)
# #             return df.to_markdown(index=False)
# #         except Exception as e:
# #             logger.error(f"Error processing structured data {file_path}: {e}")
# #             return ""

# #     def process_text(self, file_path: Path) -> str:
# #         """Reads plain text files."""
# #         try:
# #             return file_path.read_text(encoding='utf-8', errors='ignore')
# #         except Exception as e:
# #             logger.error(f"Error processing text file {file_path}: {e}")
# #             return ""

# #     def digitize(self, file_path: Path) -> Dict[str, Any]:
# #         """
# #         Main pipeline method to digitize a document.
# #         """
# #         logger.info(f"Digitizing document: {file_path.name}")
# #         doc_type = self.detect_file_type(file_path)
        
# #         raw_text = ""
# #         if doc_type == "PDF":
# #             raw_text = self.process_pdf(file_path)
# #         elif doc_type == "DOCX":
# #             raw_text = self.process_docx(file_path)
# #         elif doc_type == "STRUCTURED":
# #             raw_text = self.process_structured(file_path)
# #         elif doc_type == "TEXT":
# #             raw_text = self.process_text(file_path)
# #         else:
# #             logger.warning(f"Unsupported file type for {file_path.name}")
# #             return {"status": "error", "reason": "unsupported_type", "file": file_path.name}
        
# #         if not raw_text.strip():
# #             logger.warning(f"No text extracted from {file_path.name}")
# #             return {"status": "error", "reason": "empty_content", "file": file_path.name}
            
# #         logger.info(f"Successfully digitized {file_path.name} ({len(raw_text)} characters)")
# #         return {
# #             "status": "success",
# #             "file_name": file_path.name,
# #             "document_type": doc_type,
# #             "text_length": len(raw_text),
# #             "content": raw_text
# #         }

# # def document_digitization_pipeline(file_path: Path | str) -> Dict[str, Any]:
# #     """
# #     Entry point for digitizing a document.
# #     """
# #     if isinstance(file_path, str):
# #         file_path = Path(file_path)
    
# #     digitizer = DocumentDigitizer()
# #     return digitizer.digitize(file_path)

# # if __name__ == "__main__":
# #     # --- EXAMPLE USAGE & TESTING ---
# #     test_dir = Path("test_data")
# #     test_dir.mkdir(exist_ok=True)
    
# #     # 1. Test Unstructured Text
# #     txt_file = test_dir / "test_doc.txt"
# #     txt_file.write_text("Client Invoice Report - 2023-10-27\nTotal due amount is $1,250.")
# #     print("\n--- Testing TXT ---")
# #     result_txt = document_digitization_pipeline(txt_file)
# #     print(f"Status: {result_txt.get('status')}, Extracted Length: {result_txt.get('text_length')}")
    
# #     # 2. Test Structured Data (CSV)
# #     csv_file = test_dir / "test_data.csv"
# #     dummy_df = pd.DataFrame({'Product': ['Laptop', 'Mouse'], 'Quantity': [10, 50], 'Price': [1200.0, 25.0]})
# #     dummy_df.to_csv(csv_file, index=False)
# #     print("\n--- Testing CSV ---")
# #     result_csv = document_digitization_pipeline(csv_file)
# #     print(f"Status: {result_csv.get('status')}, Extracted Length: {result_csv.get('text_length')}")
# #     print("\nPreview of extracted CSV markdown:")
# #     print(result_csv.get('content'))

# import os
# from typing import List, Dict, Any, Optional

# class DocumentProcessor:
#     """
#     A class responsible for processing various document types to extract structured data.
    
#     This class encapsulates the logic for document loading, parsing, and data extraction.
#     """
    
#     def __init__(self):
#         """Initializes the processor with any necessary resources or configurations."""
#         print("DocumentProcessor initialized.")
#         # In a real-world scenario, you might initialize ML models or API clients here.
#         pass

#     def load_document(self, file_path: str) -> Optional[bytes]:
#         """
#         Loads the content of a file.
        
#         Args:
#             file_path: The full path to the document file.
            
#         Returns:
#             The raw binary content of the file, or None if loading fails.
#         """
#         if not os.path.exists(file_path):
#             print(f"Error: File not found at path: {file_path}")
#             return None
        
#         try:
#             with open(file_path, 'rb') as f:
#                 return f.read()
#         except IOError as e:
#             print(f"Error reading file {file_path}: {e}")
#             return None

#     def process_document(self, file_path: str) -> Optional[Dict[str, Any]]:
#         """
#         Main entry point to process a document.
        
#         Args:
#             file_path: The path to the document.
            
#         Returns:
#             A dictionary containing the extracted data, or None if processing fails.
#         """
#         raw_data = self.load_document(file_path)
#         if raw_data is None:
#             return None
        
#         # Determine processing method based on file extension
#         extension = os.path.splitext(file_path)[1].lower()
        
#         try:
#             if extension in ['.pdf']:
#                 print("Attempting PDF processing...")
#                 # Placeholder for actual PDF parsing logic (requires libraries like PyPDF2, pdfminer)
#                 extracted_data = self._process_pdf(raw_data)
#             elif extension in ['.txt']:
#                 print("Attempting TXT processing...")
#                 extracted_data = self._process_text(raw_data)
#             elif extension in ['.docx']:
#                 print("Attempting DOCX processing...")
#                 # Placeholder for actual DOCX parsing logic (requires libraries like python-docx)
#                 extracted_data = self._process_docx(raw_data)
#             else:
#                 print(f"Warning: Unsupported file type: {extension}")
#                 return None
            
#             return {
#                 "source_file": file_path,
#                 "status": "SUCCESS",
#                 "extracted_data": extracted_data
#             }
#         except Exception as e:
#             print(f"Critical error during processing of {file_path}: {e}")
#             return {
#                 "source_file": file_path,
#                 "status": "FAILED",
#                 "error_message": str(e)
#             }

#     # --- Private Parsing Methods (Separation of Concerns) ---

#     def _process_pdf(self, raw_data: bytes) -> Dict[str, Any]:
#         """Simulates complex PDF parsing."""
#         # TODO: Implement actual PDF extraction logic here
#         return {"text_content": f"[PDF content extracted from {len(raw_data)} bytes]", "metadata": {"pages": 1}}

#     def _process_text(self, raw_data: bytes) -> Dict[str, Any]:
#         """Processes plain text files."""
#         try:
#             text_content = raw_data.decode('utf-8')
#             # Simple extraction logic for demonstration
#             lines = text_content.strip().split('\n')
#             return {"text_content": text_content, "line_count": len(lines)}
#         except UnicodeDecodeError:
#             raise ValueError("Could not decode file as UTF-8.")

#     def _process_docx(self, raw_data: bytes) -> Dict[str, Any]:
#         """Simulates DOCX parsing."""
#         # TODO: Implement actual DOCX extraction logic here
#         return {"text_content": f"[DOCX content extracted from {len(raw_data)} bytes]", "metadata": {"sections": 1}}

# # --- Example Usage ---

# if __name__ == "__main__":
#     # --- Setup Mock Files for Testing ---
#     # Create dummy files to test the structure
#     MOCK_PDF_PATH = "test_document.pdf"
#     MOCK_TXT_PATH = "test_document.txt"
#     MOCK_DOCX_PATH = "test_document.docx"
    
#     try:
#         with open(MOCK_TXT_PATH, 'w', encoding='utf-8') as f:
#             f.write("This is a test document.\nLine two contains important data.\nEnd.")
#         # Create empty dummy files for other types to test file existence checks
#         with open(MOCK_PDF_PATH, 'wb') as f:
#             f.write(b'%PDF-1.4')
#         with open(MOCK_DOCX_PATH, 'wb') as f:
#             f.write(b'PK\x03\x04')
            
#         print("\n" + "="*50)
#         print("STARTING DOCUMENT PROCESSING DEMO")
#         print("="*50)

#         # 1. Test TXT file
#         processor = DocumentProcessor()
#         txt_result = processor.process_document(MOCK_TXT_PATH)
#         print("\n--- TXT Result ---")
#         print(txt_result)

#         # 2. Test PDF file (Mocked)
#         pdf_result = processor.process_document(MOCK_PDF_PATH)
#         print("\n--- PDF Result ---")
#         print(pdf_result)

#         # 3. Test non-existent file
#         print("\n--- Missing File Test ---")
#         missing_result = processor.process_document("non_existent_file.xyz")
#         print(missing_result)

#     finally:
#         # Cleanup mock files
#         for path in [MOCK_TXT_PATH, MOCK_PDF_PATH, MOCK_DOCX_PATH]:
#             if os.path.exists(path):
#                 os.remove(path)
#         print("\nCleanup complete.")

import logging
from pathlib import Path
from typing import Dict, Any, List, Optional

import fitz
import docx
import pandas as pd
from PIL import Image
import pytesseract

from llama_index import Document, GPTVectorStoreIndex, LLMPredictor, ServiceContext, PromptHelper
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.llms import OpenAI

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- 1. File Loader/Ingestor Module ---
class DocumentTextExtractor:
    """Extracts text from document files so they can be indexed."""

    def extract(self, file_path: Path) -> Dict[str, Any]:
        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self._extract_pdf(file_path)
        if extension == '.docx':
            return self._extract_docx(file_path)
        if extension in ['.txt', '.md']:
            return self._extract_text(file_path)
        if extension == '.csv':
            return self._extract_csv(file_path)
        if extension in ['.xls', '.xlsx']:
            return self._extract_excel(file_path)
        if extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return self._extract_image(file_path)

        raise ValueError(f'Unsupported document type: {extension}')

    def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        logger.info('Extracting PDF text from %s', file_path)
        text_blocks: List[str] = []

        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    page_text = page.get_text().strip()
                    if page_text:
                        text_blocks.append(page_text)
        except Exception as exc:
            logger.warning('PDF extraction failed for %s: %s', file_path, exc)
            return self._extract_image(file_path)

        return self._build_document_payload(file_path, '\n\n'.join(text_blocks))

    def _extract_docx(self, file_path: Path) -> Dict[str, Any]:
        logger.info('Extracting DOCX text from %s', file_path)
        paragraphs: List[str] = []

        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())

        return self._build_document_payload(file_path, '\n\n'.join(paragraphs))

    def _extract_text(self, file_path: Path) -> Dict[str, Any]:
        logger.info('Extracting plain text from %s', file_path)
        text = file_path.read_text(encoding='utf-8', errors='ignore')
        return self._build_document_payload(file_path, text)

    def _extract_csv(self, file_path: Path) -> Dict[str, Any]:
        logger.info('Extracting CSV text from %s', file_path)
        df = pd.read_csv(file_path)
        return self._build_document_payload(file_path, df.to_markdown(index=False))

    def _extract_excel(self, file_path: Path) -> Dict[str, Any]:
        logger.info('Extracting Excel text from %s', file_path)
        df = pd.read_excel(file_path)
        return self._build_document_payload(file_path, df.to_markdown(index=False))

    def _extract_image(self, file_path: Path) -> Dict[str, Any]:
        logger.info('Extracting image text from %s using OCR', file_path)
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return self._build_document_payload(file_path, text)

    def _build_document_payload(self, file_path: Path, text: str) -> Dict[str, Any]:
        return {
            'source': str(file_path),
            'text': text.strip(),
            'metadata': {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_type': file_path.suffix.lower()
            }
        }


# --- 2. Text Extractor & Chunker Module ---
class TextProcessor:
    """Splits raw text into manageable chunks."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_text(self, raw_text: str) -> List[str]:
        """Splits text into overlapping chunks."""
        logger.info('Chunking text into %d character segments with %d overlap', self.chunk_size, self.chunk_overlap)
        chunks = []
        start = 0

        while start < len(raw_text):
            end = start + self.chunk_size
            chunks.append(raw_text[start:end].strip())
            start += self.chunk_size - self.chunk_overlap

        return [chunk for chunk in chunks if chunk]


class LlamaIndexIngestor:
    """Builds and queries a llamaindex vector index for documents."""

    def __init__(self, embedding_model: str = 'text-embedding-ada-002', llm_model: str = 'gpt-3.5-turbo'):
        self.embedding_model = embedding_model
        self.llm_model = llm_model
        self.service_context = self._build_service_context()
        self.index: Optional[GPTVectorStoreIndex] = None

    def _build_service_context(self) -> ServiceContext:
        llm_predictor = LLMPredictor(llm=OpenAI(model_name=self.llm_model, temperature=0))
        prompt_helper = PromptHelper(max_input_size=4096, num_output=512, max_chunk_overlap=150)
        embed_model = OpenAIEmbedding(model=self.embedding_model)

        return ServiceContext.from_defaults(
            llm_predictor=llm_predictor,
            embed_model=embed_model,
            prompt_helper=prompt_helper
        )

    def build_index(self, documents: List[Document], persist_dir: Optional[Path] = None) -> GPTVectorStoreIndex:
        logger.info('Building llamaindex vector index for %d documents', len(documents))
        self.index = GPTVectorStoreIndex.from_documents(documents, service_context=self.service_context)

        if persist_dir:
            persist_dir.mkdir(parents=True, exist_ok=True)
            self.index.storage_context.persist(persist_dir)
            logger.info('Index persisted to %s', persist_dir)

        return self.index

    def load_index(self, persist_dir: Path) -> GPTVectorStoreIndex:
        logger.info('Loading persisted index from %s', persist_dir)
        self.index = GPTVectorStoreIndex.load_from_disk(persist_dir)
        return self.index

    def query(self, query_text: str, top_k: int = 5) -> str:
        if self.index is None:
            raise RuntimeError('Index has not been built or loaded yet.')

        query_engine = self.index.as_query_engine(similarity_top_k=top_k)
        response = query_engine.query(query_text)
        return str(response)


# --- 4. Orchestrator (The Main Pipeline) ---
class DocumentDigitizationPipeline:
    """Pipeline that extracts document text, builds a llamaindex, and supports question answering."""

    def __init__(self, index_dir: Optional[str] = None):
        self.extractor = DocumentTextExtractor()
        self.ingestor = LlamaIndexIngestor()
        self.index_dir = Path(index_dir) if index_dir else None

    def prepare_documents(self, file_paths: List[Path]) -> List[Document]:
        documents: List[Document] = []

        for file_path in file_paths:
            logger.info('Preparing %s', file_path)
            extracted = self.extractor.extract(file_path)
            if not extracted['text']:
                logger.warning('No text extracted from %s', file_path)
                continue

            documents.append(Document(text=extracted['text'], metadata=extracted['metadata']))

        return documents

    def ingest_files(self, file_paths: List[Path]) -> GPTVectorStoreIndex:
        documents = self.prepare_documents(file_paths)
        if not documents:
            raise ValueError('No documents were extracted successfully.')

        return self.ingestor.build_index(documents, persist_dir=self.index_dir)

    def query(self, query_text: str, top_k: int = 5) -> str:
        return self.ingestor.query(query_text, top_k=top_k)


# =================================================================
#                           EXAMPLE USAGE
# =================================================================

def get_supported_files(folder: Path) -> List[Path]:
    supported = {'.pdf', '.docx', '.txt', '.md', '.csv', '.xls', '.xlsx', '.jpg', '.jpeg', '.png', '.tiff', '.bmp'}
    return [path for path in folder.iterdir() if path.suffix.lower() in supported]


if __name__ == '__main__':
    data_folder = Path('documents')
    data_folder.mkdir(exist_ok=True)

    logger.info('Starting document digitization and llamaindex ingestion pipeline')

    file_list = get_supported_files(data_folder)
    if not file_list:
        logger.warning('No supported documents found in %s', data_folder)
    else:
        pipeline = DocumentDigitizationPipeline(index_dir=Path('index_store'))
        index = pipeline.ingest_files(file_list)

        logger.info('Document ingestion complete. Indexed %d documents.', len(file_list))

        sample_query = 'What is the main topic of the documents?'
        answer = pipeline.query(sample_query)
        print('--- Sample Query ---')
        print(sample_query)
        print(answer)